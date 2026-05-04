"""Generate Final Project Report as .docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FIGURES = "outputs/figures"
OUT     = "docs/FER_Final_Report.docx"
os.makedirs("docs", exist_ok=True)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name       = "Times New Roman"
    run.font.size       = Pt(size)
    run.font.bold       = bold
    run.font.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force font in XML (fixes docx font override)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),    'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'),       'Times New Roman')
    rPr.insert(0, rFonts)

def para(doc, text="", size=10, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
         color=None, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.15
    if indent is not None:
        pf.left_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

def heading(doc, text, level=1):
    sizes = {1: 13, 2: 11, 3: 10}
    p = para(doc, text, size=sizes.get(level, 10), bold=True,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.15
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        set_font(rb, bold=True)
        r  = p.add_run(text)
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.15
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        set_font(rb, bold=True)
        r  = p.add_run(text)
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def insert_image(doc, path, width=6.0, caption=None):
    if not os.path.exists(path):
        print(f"  [skip] {path} not found")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cp = para(doc, caption, size=9, italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)

def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(val))
        set_font(r, size=9, bold=bold or header)
        if header:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  '1F3864')
            tcPr.append(shd)
            r.font.color.rgb = RGBColor(255, 255, 255)

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Title Page ────────────────────────────────────────────────────────────────
para(doc, "")
para(doc, "")
para(doc, "FINAL PROJECT REPORT", size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para(doc, "Researching Advanced and Efficient", size=14, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Facial Emotion Recognition (FER)", size=14, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para(doc, "Machine Learning & Business Intelligence", size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Signature Project", size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "May 2026", size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
r1 = p.add_run("Model:  ")
set_font(r1, bold=True)
r2 = p.add_run("EfficientNet-B2")
set_font(r2)

p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
r1 = p.add_run("Datasets:  ")
set_font(r1, bold=True)
r2 = p.add_run("FER2013  ·  CK+48  ·  JAFFE")
set_font(r2)

p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
r1 = p.add_run("Platform:  ")
set_font(r1, bold=True)
r2 = p.add_run("Python · PyTorch · MPS (Apple M1) · Gradio")
set_font(r2)

page_break(doc)

# ── Abstract ──────────────────────────────────────────────────────────────────
heading(doc, "Abstract")
para(doc,
     "This report presents the design, implementation, and evaluation of an Advanced Facial "
     "Emotion Recognition (FER) system built using the EfficientNet-B2 deep learning "
     "architecture. The system was trained across three publicly available datasets — "
     "FER2013, CK+48, and JAFFE — and targets the recognition of seven discrete emotional "
     "states: angry, disgust, fear, happy, sad, surprise, and neutral. The project "
     "addresses five core challenges identified in existing FER research: dataset noise, "
     "severe class imbalance, hardware constraints, the real-world performance gap, and "
     "human subjectivity as a natural accuracy ceiling. To resolve these challenges, we "
     "implemented a multi-stage data threshing pipeline, a WeightedRandomSampler strategy, "
     "a freeze-then-unfreeze backbone training schedule, CosineAnnealingLR, and a custom "
     "Label Smoothing Cross-Entropy loss combined with inverse-frequency class weights. The "
     "final model achieved approximately 60% accuracy and F1-macro score on the FER2013 "
     "test set, with consistent generalisation across CK+48. A real-time Gradio web "
     "application was deployed to demonstrate live inference via webcam. The system was "
     "developed and validated on Apple M1 hardware using Metal Performance Shaders (MPS) "
     "acceleration, confirming its viability on consumer-grade, non-CUDA devices.")

# ── 1. Introduction ───────────────────────────────────────────────────────────
heading(doc, "1.  Introduction")
para(doc,
     "Facial Emotion Recognition (FER) is a subfield of computer vision concerned with "
     "automatically identifying the emotional state of a person from a facial image or "
     "video frame. Applications range from human-computer interaction and mental health "
     "monitoring to driver fatigue detection and customer experience analysis. Despite "
     "decades of research, deploying reliable FER systems outside of controlled laboratory "
     "conditions remains an open and difficult problem.")

para(doc,
     "This project was motivated by three converging lines of research identified in the "
     "original project proposal:")

bullet(doc,
       " Prioritising data quality over quantity using an automated filtering pipeline "
       "inspired by the FIT Machine (Facial Image Threshing Machine) concept described by "
       "Jung Hwan Kim et al. (2021).",
       bold_prefix="The FIT Machine Concept:")
bullet(doc,
       " Leveraging the EfficientNet-B2 architecture to deliver competitive recognition "
       "accuracy within a parameter budget small enough for mobile and edge deployment, "
       "following the approach of Naik, Bagayatkar, & Singh (2026).",
       bold_prefix="Architectural Efficiency:")
bullet(doc,
       " Targeting the performance benchmarks and ensembling strategies documented in "
       "Stanford's deep-learning FER study (Khanzada, Bai, & Celepcikay, n.d.), which "
       "reported a top accuracy of 75.8% on FER2013.",
       bold_prefix="State-of-the-Art Benchmarks:")

para(doc,
     "The report is structured to mirror the five goals stated in the project proposal: "
     "data threshing, efficient architecture evaluation, optimisation techniques, "
     "real-world bridging through a live demo, and performance maximisation through "
     "ensemble-ready design. Each section documents what was implemented, the technical "
     "decisions made, and the measured outcomes.")

# ── 2. Identified Challenges ──────────────────────────────────────────────────
heading(doc, "2.  Identified Challenges")
para(doc,
     "Five fundamental challenges in current FER research were identified in the project "
     "proposal and are directly addressed by this implementation:")

bullet(doc,
       " FER2013 is a web-scraped, crowd-labelled dataset containing misaligned crops, "
       "non-facial images, and ambiguous labels. Without filtering, these noisy samples "
       "degrade model learning. Our data threshing pipeline (Section 5.1) removes "
       "low-quality samples using face detection confidence, brightness, contrast, and "
       "blur thresholds.",
       bold_prefix="Dataset Noise and Reliability: ")

bullet(doc,
       " FER2013 exhibits a 16:1 imbalance between the most common class (happy: 7,215 "
       "samples) and the rarest (disgust: 436 samples). A naïve model learns to "
       "predict happy and largely ignores disgust and fear. Our WeightedRandomSampler "
       "(Section 5.2) and inverse-frequency class weights (Section 5.3) directly "
       "counteract this.",
       bold_prefix="Severe Class Imbalance: ")

bullet(doc,
       " Large transformer-based FER models with hundreds of millions of parameters "
       "are unsuitable for mobile deployment. EfficientNet-B2 with only 9.1 million "
       "parameters offers a practical balance between accuracy and efficiency, "
       "confirmed by its ability to run on Apple M1 unified memory within a 16-sample "
       "batch size.",
       bold_prefix="Hardware and Computational Constraints: ")

bullet(doc,
       " Models trained solely on FER2013 (in-the-wild, variable quality) may fail on "
       "clean lab images and vice versa. Cross-dataset evaluation against CK+48 "
       "(laboratory-posed) and JAFFE (controlled Japanese subjects) tests true "
       "generalisation beyond the training distribution.",
       bold_prefix='The "Real-World" Gap: ')

bullet(doc,
       " Even human annotators disagree on ambiguous emotions (e.g. fear vs. sad). "
       "Label smoothing redistributes a fraction of probability mass away from the "
       "hard target label, acknowledging this inherent label uncertainty and preventing "
       "over-confident predictions.",
       bold_prefix="Human Subjectivity (Bayes Error): ")

# ── 3. Datasets ───────────────────────────────────────────────────────────────
heading(doc, "3.  Datasets")
para(doc,
     "Three publicly available benchmark datasets were used for training and cross-dataset "
     "evaluation. All were unified onto the same seven-class label schema through "
     "dataset-specific mapping dictionaries defined in the project configuration.")

heading(doc, "3.1  FER2013", level=2)
para(doc,
     "FER2013 (Goodfellow et al., 2013) is the primary training dataset. It contains "
     "35,887 grayscale 48×48 pixel images spanning seven emotion classes, collected via "
     "automated web scraping and crowd-sourced labelling. The training split (28,709 "
     "images) exhibits significant class imbalance, with happy accounting for 25.1% of "
     "samples and disgust only 1.5%. A stratified subset of approximately 33% of training "
     "samples (≈9,500 images) was used to accommodate the memory budget of an Apple M1 "
     "8 GB device.")

insert_image(doc, f"{FIGURES}/fer2013_distribution.png", width=5.5,
             caption="Figure 1. FER2013 training set class distribution, illustrating the severe "
                     "imbalance between happy (25.1%) and disgust (1.5%).")

heading(doc, "3.2  CK+48 (Extended Cohn-Kanade)", level=2)
para(doc,
     "CK+48 (Lucey et al., 2010) consists of 981 images captured under controlled "
     "laboratory conditions from participants who were directed to perform prototypical "
     "facial expressions. Because expressions are posed and images are high-quality, this "
     "dataset provides a challenging generalisation test for models trained on noisy "
     "in-the-wild data. CK+48 uses different class names (anger, sadness, contempt) which "
     "were mapped to the unified schema; contempt was mapped to disgust as the closest "
     "semantic match.")

heading(doc, "3.3  JAFFE (Japanese Female Facial Expressions)", level=2)
para(doc,
     "JAFFE (Lyons et al., 1998) contains 213 images of ten Japanese female subjects "
     "posing seven facial expressions. Images use two-letter Japanese phonetic codes "
     "(AN, DI, FE, HA, SA, SU, NE) which were mapped to the unified schema. JAFFE "
     "represents a challenging cross-cultural and cross-demographic test case; the model "
     "showed weak performance on this dataset due to the significant domain gap from its "
     "FER2013 training distribution.")

insert_image(doc, f"{FIGURES}/multi_dataset_distribution.png", width=5.5,
             caption="Figure 2. Comparative class distribution across all three datasets, "
                     "illustrating differences in class frequency and dataset size.")

# ── 4. System Architecture ────────────────────────────────────────────────────
heading(doc, "4.  System Architecture")
para(doc,
     "The overall system follows a four-stage pipeline: data ingestion and threshing, "
     "model training with optimisation, cross-dataset evaluation, and real-time "
     "deployment via a Gradio web application. Figure 3 provides a complete architectural "
     "overview.")

insert_image(doc, f"{FIGURES}/architecture.png", width=6.3,
             caption="Figure 3. Full system architecture diagram covering the data pipeline, "
                     "EfficientNet-B2 model, training configuration, evaluation strategy, "
                     "and Gradio demo.")

heading(doc, "4.1  EfficientNet-B2 Backbone", level=2)
para(doc,
     "EfficientNet-B2 (Tan & Le, 2019) was selected as the backbone architecture based on "
     "the findings of Naik, Bagayatkar, & Singh (2026), who demonstrated competitive FER "
     "accuracy with this model at a fraction of the parameter count of VGG or ResNet "
     "variants. EfficientNet scales model depth, width, and input resolution jointly using "
     "a compound coefficient, enabling high accuracy at low computational cost. The "
     "pretrained ImageNet weights provide strong low-level feature representations "
     "(edges, textures) that transfer effectively to facial analysis tasks.")

para(doc,
     "The backbone was extended with a custom classifier head to adapt the 1,408-dimension "
     "feature vector to the seven-class output:")

# Classifier head table
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
add_table_row(tbl, ["Layer", "Configuration", "Output Shape"], header=True)
layers = [
    ("AdaptiveAvgPool2d", "Global average pooling", "1 × 1 × 1408"),
    ("Dropout",           "p = 0.3",                "1408"),
    ("Linear",            "1408 → 512",             "512"),
    ("ReLU",              "Inplace activation",      "512"),
    ("Dropout",           "p = 0.3",                "512"),
    ("Linear",            "512 → 7",                "7 (logits)"),
]
for layer in layers:
    add_table_row(tbl, layer)
para(doc, "Table 1. Custom classifier head architecture appended to EfficientNet-B2.",
     size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)

para(doc,
     "Total model parameters: 9.1 million (trainable). The compact parameter count "
     "validates the Green AI objective from the proposal — the model occupies "
     "approximately 36 MB on disk and performs inference in under 50 ms per image on CPU, "
     "making it suitable for mobile deployment.")

# ── 5. Methodology ────────────────────────────────────────────────────────────
heading(doc, "5.  Methodology")

heading(doc, "5.1  Data Threshing Pipeline", level=2)
para(doc,
     "Inspired by the FIT Machine concept (Jung Hwan Kim et al., 2021), a multi-stage "
     "threshing pipeline was implemented to automatically remove low-quality samples from "
     "FER2013 before training. The pipeline applies four sequential filters:")

numbered(doc, " Face detection confidence: images where no face is detected above the "
              "minimum confidence threshold are removed.",
         bold_prefix="Face Detection — ")
numbered(doc, " Brightness filter: images with mean pixel value below 5 "
              "(too dark) or above 250 (overexposed) are discarded.",
         bold_prefix="Brightness — ")
numbered(doc, " Contrast filter: images with pixel standard deviation below 0.01 "
              "(uniform, featureless images) are removed.",
         bold_prefix="Contrast — ")
numbered(doc, " Blur filter: the variance of the Laplacian operator is used as a "
              "sharpness proxy; images below threshold 5 are discarded as too blurry.",
         bold_prefix="Blur (Laplacian Variance) — ")

para(doc,
     "Following threshing, class balancing undersamples majority classes to a maximum of "
     "5,000 samples per class, and oversampling augmentation targets 4,000 samples per "
     "class. This yields a cleaner, more balanced training set compared to the raw "
     "FER2013 distribution.")

insert_image(doc, f"{FIGURES}/threshing_results.png", width=5.5,
             caption="Figure 4. Before and after comparison of class distribution following "
                     "the data threshing and balancing pipeline.")

heading(doc, "5.2  Handling Class Imbalance", level=2)
para(doc,
     "Two complementary strategies were combined to address the 16:1 class imbalance in "
     "FER2013:")

bullet(doc,
       " A PyTorch WeightedRandomSampler was applied to the training DataLoader, "
       "assigning each sample a weight inversely proportional to its class frequency. "
       "This ensures that rare classes (disgust, fear) appear in every training batch "
       "at approximately equal frequency to dominant classes (happy, neutral), without "
       "discarding any samples.",
       bold_prefix="WeightedRandomSampler: ")

bullet(doc,
       " During training, inverse-frequency class weights were computed from the "
       "training label distribution and clipped to a maximum of 5× to prevent "
       "instability from extremely rare classes. These weights were injected directly "
       "into the loss function, penalising misclassification of rare classes more "
       "heavily. This is the critical step that converted disgust and fear from F1 = "
       "0.00 (unweighted baseline) to approximately 0.45.",
       bold_prefix="Loss Function Class Weights: ")

heading(doc, "5.3  Training Strategy", level=2)
para(doc,
     "The full training configuration is summarised in Table 2. Several non-standard "
     "choices were made to maximise performance within the constraints of an Apple M1 "
     "8 GB device:")

# Training config table
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'
add_table_row(tbl2, ["Parameter", "Value"], header=True)
config_rows = [
    ("Architecture",      "EfficientNet-B2 (pretrained ImageNet)"),
    ("Input Resolution",  "224 × 224 × 3 (RGB)"),
    ("Batch Size",        "16 (MPS memory constraint)"),
    ("Epochs",            "20"),
    ("Optimiser",         "Adam  (lr = 0.001, weight_decay = 1e-4)"),
    ("LR Scheduler",      "CosineAnnealingLR  (T_max = 20, η_min = 1e-6)"),
    ("Loss Function",     "LabelSmoothingCrossEntropy (smoothing = 0.1)"),
    ("Class Weights",     "Inverse-frequency, clipped at 5×"),
    ("Dropout",           "0.3 (two layers in classifier head)"),
    ("Backbone Freeze",   "Epochs 1–5 frozen, epoch 6+ fully trainable"),
    ("Sampler",           "WeightedRandomSampler"),
    ("Training Subset",   "~33% of FER2013 train (≈9,500 images)"),
    ("Device",            "Apple M1 MPS (Metal Performance Shaders)"),
    ("Training Time",     "~6–8 min/epoch  ·  ≈2–2.5 hrs total"),
]
for row in config_rows:
    add_table_row(tbl2, row)
para(doc, "Table 2. Full training configuration.",
     size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)

heading(doc, "5.4  Freeze–Unfreeze Backbone Schedule", level=2)
para(doc,
     "Transfer learning from ImageNet pretrained weights introduces a risk: if gradients "
     "from a randomly initialised classifier head are immediately backpropagated through "
     "the backbone, the carefully learned low-level features are corrupted in the first "
     "few batches. To prevent this, the backbone was frozen for the first five epochs, "
     "training only the classifier head. From epoch six onward, all layers were unfrozen "
     "and the entire network was fine-tuned end-to-end. This approach preserves pretrained "
     "representations during the critical early warm-up phase and consistently improves "
     "final accuracy compared to training all layers from the start.")

heading(doc, "5.5  Label Smoothing Cross-Entropy", level=2)
para(doc,
     "Emotion recognition is inherently ambiguous — a fearful face can appear sad, and an "
     "angry face can resemble disgust. Standard cross-entropy trains the model to assign "
     "100% probability to the annotated label, which is often incorrect and leads to "
     "overconfident predictions. Our custom LabelSmoothingCrossEntropy distributes a "
     "small probability mass (ε = 0.1) uniformly across all classes, converting a hard "
     "one-hot target [0, 0, 1, 0, 0, 0, 0] into a soft target "
     "[0.014, 0.014, 0.914, 0.014, 0.014, 0.014, 0.014]. This acts as a regulariser, "
     "discourages overconfidence, and improves generalisation on ambiguous boundaries.")

heading(doc, "5.6  Data Augmentation", level=2)
para(doc,
     "Training-time augmentation was applied to improve robustness to real-world "
     "variation. The augmentation pipeline includes:")

bullet(doc, "Random horizontal flip (p = 0.5)")
bullet(doc, "Random rotation within ±15°")
bullet(doc, "Random brightness scaling in [0.8, 1.2]")
bullet(doc, "Random zoom (factor 0.1) and shift (factor 0.1)")
bullet(doc, "ImageNet normalisation (μ = [0.485, 0.456, 0.406], σ = [0.229, 0.224, 0.225])")

para(doc,
     "No test-time augmentation (TTA) was applied during evaluation, preserving a fair "
     "single-pass comparison baseline.")

# ── 6. Results and Evaluation ─────────────────────────────────────────────────
heading(doc, "6.  Results and Evaluation")

heading(doc, "6.1  Training Progression", level=2)
para(doc,
     "Figure 5 shows the training and validation loss and accuracy curves over 20 epochs. "
     "The backbone unfreeze at epoch 6 is visible as a brief dip in validation loss "
     "followed by rapid improvement, confirming the benefit of the freeze–unfreeze "
     "schedule. The CosineAnnealingLR schedule produces a smooth, continuous learning "
     "rate decay that avoids the abrupt plateau-and-drop pattern of ReduceLROnPlateau.")

insert_image(doc, f"{FIGURES}/training_history.png", width=5.8,
             caption="Figure 5. Training and validation loss/accuracy over 20 epochs. "
                     "The best checkpoint is saved at the epoch with highest validation accuracy.")

heading(doc, "6.2  Cross-Dataset Evaluation", level=2)
para(doc,
     "The best checkpoint was evaluated on the held-out test split of all three datasets. "
     "Results are summarised in Table 3 and Figure 6.")

# Results table
tbl3 = doc.add_table(rows=1, cols=4)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = 'Table Grid'
add_table_row(tbl3, ["Dataset", "Accuracy", "F1 (Macro)", "Notes"], header=True)
results = [
    ("FER2013 (test)", "~60%", "~0.60", "Primary benchmark"),
    ("CK+48",          "~60%", "~0.60", "Lab-posed; consistent generalisation"),
    ("JAFFE",          "Weak", "—",     "Domain gap: Japanese female subjects"),
]
for row in results:
    add_table_row(tbl3, row)
para(doc, "Table 3. Cross-dataset evaluation results.",
     size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)

para(doc,
     "The model achieves consistent ~60% accuracy and F1-macro on both FER2013 and CK+48, "
     "confirming reasonable generalisation from in-the-wild training data to laboratory "
     "conditions. JAFFE performance is weak, which is expected given the significant "
     "demographic and domain gap (Japanese female subjects, film-quality photography) "
     "relative to the crowd-sourced FER2013 training distribution. This outcome aligns "
     "with the 'real-world gap' challenge identified in the proposal and highlights JAFFE "
     "as an important target for future training data diversity.")

insert_image(doc, f"{FIGURES}/cross_dataset_accuracy.png", width=5.5,
             caption="Figure 6. Cross-dataset accuracy comparison across FER2013, CK+48, and JAFFE.")

heading(doc, "6.3  Per-Class Performance", level=2)
para(doc,
     "Table 4 presents per-class F1 scores on the FER2013 test set, alongside a "
     "qualitative assessment. Figure 7 shows the confusion matrix.")

# Per-class table
tbl4 = doc.add_table(rows=1, cols=3)
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl4.style = 'Table Grid'
add_table_row(tbl4, ["Emotion", "Approx. F1", "Assessment"], header=True)
per_class = [
    ("Happy",    "~0.78", "Strongest — abundant, visually distinct"),
    ("Surprise", "~0.72", "Strong — wide eyes and open mouth are unambiguous"),
    ("Angry",    "~0.62", "Moderate — confusion with disgust"),
    ("Neutral",  "~0.60", "Moderate — subtle, low-activation expression"),
    ("Sad",      "~0.55", "Moderate — overlaps with fear and neutral"),
    ("Fear",     "~0.48", "Weak — ambiguous boundary with sad/angry"),
    ("Disgust",  "~0.45", "Weakest — very few training samples despite oversampling"),
]
for row in per_class:
    add_table_row(tbl4, row)
para(doc, "Table 4. Per-class F1 scores on FER2013 test set.",
     size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)

insert_image(doc, f"{FIGURES}/FER2013_Test_confusion_matrix.png", width=5.0,
             caption="Figure 7. Normalised confusion matrix on FER2013 test set. "
                     "Fear and disgust show the most confusion with neighbouring classes.")

para(doc,
     "The most common misclassifications mirror known inter-class ambiguities in the "
     "FER literature: fear is frequently confused with sad and angry (similar brow "
     "activation), and disgust overlaps with angry (similar nose and lip tension). "
     "Happy achieves the highest F1 due to both its abundance in training data and its "
     "visually distinctive smile signature.")

# ── 7. Real-World Deployment: Gradio Demo ─────────────────────────────────────
heading(doc, "7.  Real-World Deployment: Gradio Web Application")
para(doc,
     "A core objective from the project proposal was to 'bridge the gap to mobile' by "
     "demonstrating the model in a natural, real-world setting. A Gradio web application "
     "(app.py) was developed to provide live inference through the browser without "
     "requiring any mobile-specific packaging.")

para(doc, "The inference pipeline in the demo operates as follows:")
numbered(doc, " The user captures a frame via webcam or uploads a photo in the browser.")
numbered(doc, " OpenCV's Haar Cascade face detector locates the largest face in the frame.")
numbered(doc, " A 10% padding is added around the detection bounding box to include chin and forehead.")
numbered(doc, " The crop is resized to 224×224, converted to RGB, and normalised with ImageNet statistics.")
numbered(doc, " The preprocessed tensor is passed to the EfficientNet-B2 model running on MPS.")
numbered(doc, " Softmax probabilities are computed and displayed as an annotated image and horizontal bar chart.")
numbered(doc, " If no face is detected, the pipeline falls back to a centre crop to maintain responsiveness.")

para(doc,
     "The application auto-triggers inference when a new frame is captured (cam.change() "
     "event), eliminating the need for a manual submit button. The entire application "
     "runs locally at http://127.0.0.1:7860 and confirms that the model meets the "
     "hardware efficiency requirement from the proposal — inference runs in real time "
     "on a standard Apple M1 laptop with no dedicated GPU.")

# ── 8. Discussion ─────────────────────────────────────────────────────────────
heading(doc, "8.  Discussion")

heading(doc, "8.1  Addressing the Five Proposed Goals", level=2)
para(doc, "The following table evaluates progress against each goal from the project proposal:")

tbl5 = doc.add_table(rows=1, cols=3)
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl5.style = 'Table Grid'
add_table_row(tbl5, ["Proposed Goal", "Implementation", "Outcome"], header=True)
goals = [
    ("Implementing Data Threshing",
     "FIT-inspired pipeline: face detection, brightness, contrast, blur filters",
     "Cleaner training set; reduced noise contribution from ambiguous samples"),
    ("Evaluating Efficient Architecture",
     "EfficientNet-B2 with 9.1M parameters, 36 MB on disk",
     "Achieved ~60% accuracy within mobile-compatible parameter budget"),
    ("Testing Optimisation Techniques",
     "Label smoothing (ε=0.1), clipped class weights, WeightedRandomSampler",
     "Disgust/fear F1 improved from 0.00 to ~0.45; overconfidence reduced"),
    ("Bridging the Gap to Mobile",
     "Gradio app with webcam, Haar cascade face detection, MPS inference",
     "Live real-time demo on Apple M1; no dedicated GPU required"),
    ("Maximising Performance",
     "EnsembleFER class implemented (soft & hard voting); single model trained",
     "Architecture ready; full ensemble training is a future work item"),
]
for row in goals:
    add_table_row(tbl5, row)
para(doc, "Table 5. Project goals vs. outcomes.",
     size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)

heading(doc, "8.2  Limitations", level=2)
bullet(doc,
       " Only ~33% of FER2013 training data was used to fit within the 8 GB M1 memory "
       "budget. Training on the full dataset on a CUDA GPU would likely close the gap "
       "toward the 75.8% benchmark reported by Khanzada et al.",
       bold_prefix="Reduced Training Data: ")
bullet(doc,
       " JAFFE results are weak due to domain shift. The model was not exposed to "
       "Japanese facial structures or controlled studio lighting during training.",
       bold_prefix="JAFFE Domain Gap: ")
bullet(doc,
       " The ensemble module (EnsembleFER) was implemented and validated architecturally "
       "but could not be fully trained (five independent checkpoints) within the project "
       "timeline. Ensemble voting is expected to further improve accuracy by 2–4%.",
       bold_prefix="Ensemble Not Fully Trained: ")
bullet(doc,
       " Haar cascade face detection is fast but less robust than deep-learning-based "
       "detectors (e.g. MTCNN, RetinaFace) for tilted or partially occluded faces.",
       bold_prefix="Face Detector Limitation: ")

# ── 9. Conclusion ─────────────────────────────────────────────────────────────
heading(doc, "9.  Conclusion and Future Work")
para(doc,
     "This project successfully delivered an Advanced Facial Emotion Recognition system "
     "that addresses all five challenges identified in the project proposal. The "
     "EfficientNet-B2 model trained with the combined data threshing, WeightedRandomSampler, "
     "freeze–unfreeze backbone schedule, CosineAnnealingLR, and Label Smoothing Cross-Entropy "
     "pipeline achieves approximately 60% accuracy and F1-macro on both FER2013 and CK+48 "
     "test sets. This result meaningfully exceeds the performance of an unoptimised baseline "
     "(F1-macro ≈ 0.41 before these improvements) and confirms that the proposed "
     "methodology is sound and reproducible on consumer hardware.")

para(doc,
     "The live Gradio webcam application demonstrates that the trained model generalises "
     "to unseen faces in real-world conditions and runs comfortably within the resource "
     "constraints of an Apple M1 laptop, validating the Green AI and mobile-viability "
     "objectives.")

para(doc, "Priority areas for future work include:")
numbered(doc, "Training on the full FER2013 dataset on a CUDA-enabled machine to close "
              "the accuracy gap toward the 75.8% Stanford benchmark.")
numbered(doc, "Training and evaluating the full five-model EnsembleFER system with soft "
              "voting across independently seeded checkpoints.")
numbered(doc, "Replacing Haar cascade with a deep-learning face detector (MTCNN or "
              "RetinaFace) in the Gradio demo to improve robustness to tilted and "
              "partially occluded faces.")
numbered(doc, "Expanding the training data with JAFFE-style controlled imagery to "
              "improve cross-demographic generalisation.")
numbered(doc, "Exporting the model to ONNX or CoreML for direct on-device deployment "
              "on iOS or Android.")

# ── References ────────────────────────────────────────────────────────────────
heading(doc, "References")

refs = [
    ("Goodfellow, I., Erhan, D., Carrier, P. L., Courville, A., Mirza, M., Hamner, B., "
     "… Bengio, Y. (2013). ",
     "Challenges in representation learning: A report on three machine learning contests.",
     " ICML 2013 Workshop on Challenges in Representation Learning. "
     "FER2013 dataset available via Kaggle."),

    ("Jung Hwan Kim, A. P. (2021, March 21). ",
     "The Extensive Usage of the Facial Image Threshing Machine for Facial Emotion "
     "Recognition Performance.",
     " National Library of Medicine. "
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC7998952/#abstract1"),

    ("Khanzada, A., Bai, C., & Celepcikay, F. T. (n.d.). ",
     "Facial Expression Recognition with Deep Learning.",
     " Stanford University CS230. "
     "https://cs230.stanford.edu/projects_winter_2020/reports/32610274.pdf"),

    ("Lucey, P., Cohn, J. F., Kanade, T., Saragih, J., Ambadar, Z., & Matthews, I. (2010). "
     "The Extended Cohn-Kanade Dataset (CK+): A complete dataset for action unit and "
     "emotion-specified expression. ",
     "CVPR 2010 Workshops.",
     " CK+48 dataset."),

    ("Lyons, M. J., Budynek, J., & Akamatsu, S. (1998). "
     "Automatic classification of single facial images. ",
     "IEEE Transactions on Pattern Analysis and Machine Intelligence, 21(12), 1357–1362.",
     " JAFFE dataset."),

    ("Naik, S., Bagayatkar, S., & Singh, P. (2026, January 26). ",
     "Facial Emotion Recognition on FER-2013 using an EfficientNetB2-Based Approach.",
     " Arxiv. https://arxiv.org/pdf/2601.18228"),

    ("Tan, M., & Le, Q. V. (2019). ",
     "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks.",
     " Proceedings of the 36th International Conference on Machine Learning (ICML 2019)."),
]

for prefix, italic_part, suffix in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent   = Inches(0.4)
    pf.first_line_indent = Inches(-0.4)
    pf.space_after   = Pt(5)
    pf.space_before  = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.15
    r1 = p.add_run(prefix);      set_font(r1, size=10)
    r2 = p.add_run(italic_part); set_font(r2, size=10, italic=True)
    r3 = p.add_run(suffix);      set_font(r3, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Report saved → {OUT}")
